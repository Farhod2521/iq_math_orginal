from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

doc = Document()
section = doc.sections[0]
section.top_margin = Cm(2)
section.bottom_margin = Cm(2)
section.left_margin = Cm(2.5)
section.right_margin = Cm(2.5)


def shade_cell(cell, hex_fill):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), hex_fill)
    tcPr.append(shd)


def shade_run(run, hex_fill):
    rPr = run._r.get_or_add_rPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), hex_fill)
    rPr.append(shd)


def h1(text):
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.bold = True
    r.font.size = Pt(15)
    r.font.color.rgb = RGBColor(0x1A, 0x73, 0xE8)
    p.paragraph_format.space_before = Pt(14)
    p.paragraph_format.space_after = Pt(4)


def h2(text):
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.bold = True
    r.font.size = Pt(12)
    r.font.color.rgb = RGBColor(0x0D, 0x47, 0xA1)
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after = Pt(3)


def endpoint_line(method, url, note=''):
    colors = {'GET': '28A745', 'POST': '0086D4', 'PUT': 'FF8C00', 'DELETE': 'DC3545'}
    p = doc.add_paragraph()
    r1 = p.add_run(f'  {method}  ')
    r1.bold = True
    r1.font.size = Pt(10)
    r1.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
    shade_run(r1, colors.get(method, '888888'))
    r2 = p.add_run(f'  {url}')
    r2.font.name = 'Courier New'
    r2.font.size = Pt(10)
    r2.bold = True
    if note:
        r3 = p.add_run(f'   ⚠ {note}')
        r3.font.size = Pt(9)
        r3.italic = True
        r3.font.color.rgb = RGBColor(0x99, 0x55, 0x00)
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after = Pt(2)


def lbl(text):
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.bold = True
    r.font.size = Pt(10)
    r.font.color.rgb = RGBColor(0x22, 0x22, 0x22)
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(1)


def code_block(text):
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.font.name = 'Courier New'
    r.font.size = Pt(9)
    r.font.color.rgb = RGBColor(0x1E, 0x1E, 0x1E)
    pPr = p._p.get_or_add_pPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), 'F0F0F0')
    pPr.append(shd)
    p.paragraph_format.left_indent = Inches(0.25)
    p.paragraph_format.space_after = Pt(6)


def info(text, color='555555'):
    p = doc.add_paragraph()
    r = p.add_run(f'  {text}')
    r.font.size = Pt(9)
    r.italic = True
    rgb = tuple(int(color[i:i+2], 16) for i in (0, 2, 4))
    r.font.color.rgb = RGBColor(*rgb)
    p.paragraph_format.space_after = Pt(3)


def tbl(headers, rows):
    t = doc.add_table(rows=1 + len(rows), cols=len(headers))
    t.style = 'Table Grid'
    hr = t.rows[0].cells
    for i, h in enumerate(headers):
        hr[i].text = h
        shade_cell(hr[i], '1A73E8')
        for para in hr[i].paragraphs:
            for run in para.runs:
                run.bold = True
                run.font.size = Pt(9)
                run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
    for ri, row_data in enumerate(rows):
        row = t.rows[ri + 1].cells
        fill = 'EEF4FF' if ri % 2 == 0 else 'FFFFFF'
        for ci, val in enumerate(row_data):
            row[ci].text = val
            shade_cell(row[ci], fill)
            for para in row[ci].paragraphs:
                for run in para.runs:
                    run.font.size = Pt(9)
    doc.add_paragraph()


def box_title(text, fill='0D47A1'):
    p = doc.add_paragraph()
    r = p.add_run(f'  {text}  ')
    r.bold = True
    r.font.size = Pt(11)
    r.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
    shade_run(r, fill)
    p.paragraph_format.space_before = Pt(10)
    p.paragraph_format.space_after = Pt(4)


# ══════════════════════════════════════════════════════
#  SARLAVHA
# ══════════════════════════════════════════════════════
t = doc.add_paragraph()
t.alignment = WD_ALIGN_PARAGRAPH.CENTER
tr = t.add_run('app_book — Kitoblar CRUD API')
tr.bold = True
tr.font.size = Pt(22)
tr.font.color.rgb = RGBColor(0x1A, 0x73, 0xE8)

s = doc.add_paragraph()
s.alignment = WD_ALIGN_PARAGRAPH.CENTER
sr = s.add_run('Online (PDF) va Oflayn (yetkazib berish) kitoblar — Frontend qo\'llanma')
sr.italic = True
sr.font.size = Pt(11)
sr.font.color.rgb = RGBColor(0x77, 0x77, 0x77)
doc.add_paragraph()

info('Base URL:  https://api.iqmath.uz/api/v1/book/')
info('Auth:      Authorization: Bearer <access_token>  — barcha so\'rovlarda talab qilinadi')
info('Fayl yuklashda: Content-Type: multipart/form-data  |  Oddiy so\'rovlarda: application/json')

# ══════════════════════════════════════════════════════
#  IKKI TUR TUSHUNTIRISH
# ══════════════════════════════════════════════════════
h1('KITOB TURLARI')

box_title('1-tur: Online PDF kitob', '0086D4')
tbl(
    ['Field', 'Qiymat', 'Izoh'],
    [
        ['is_offline', 'false  yoki  null', 'Online kitob — yetkazib berish yo\'q'],
        ['quantity',   'null',              'Ombor soni ko\'rsatilmaydi'],
        ['file',       'PDF fayl',          'Asosiy kontent — PDF fayl yuklanadi'],
        ['cover_image','Rasm (ixtiyoriy)',  'Muqova rasmi'],
    ]
)

box_title('2-tur: Oflayn (jismoniy) kitob', '28A745')
tbl(
    ['Field', 'Qiymat', 'Izoh'],
    [
        ['is_offline', 'true',         'Oflayn kitob — yetkazib berish mumkin'],
        ['quantity',   'Musbat son',   'Ombordagi nusxa soni (masalan: 50)'],
        ['file',       'null',         'PDF fayl shart emas'],
        ['cover_image','Rasm (ixtiyoriy)', 'Muqova rasmi'],
    ]
)

# ══════════════════════════════════════════════════════
#  GET LIST
# ══════════════════════════════════════════════════════
h1('1. Kitoblar ro\'yxati')
endpoint_line('GET', '/api/v1/book/books/')

lbl('Query parametrlar (ixtiyoriy):')
tbl(
    ['Param', 'Misol', 'Izoh'],
    [
        ['category',    '?category=1',      'Faqat shu kategoriya kitoblari'],
        ['status',      '?status=active',   'active  |  inactive  |  draft'],
        ['tag',         '?tag=2',           'Faqat shu teg biriktirilgan kitoblar'],
    ]
)

lbl('Response  200 OK:')
code_block(
    '[\n'
    '  {\n'
    '    "id": 1,\n'
    '    "name": "Algebra asoslari",\n'
    '    "description": "To\'liq qo\'llanma",\n'
    '    "category": { "id": 1, "name": "Matematika", "is_active": true, "created_at": "..." },\n'
    '    "tags": [ { "id": 1, "name": "beginner", "is_active": true, "created_at": "..." } ],\n'
    '    "file": "https://api.iqmath.uz/media/books/files/algebra.pdf",\n'
    '    "cover_image": "https://api.iqmath.uz/media/books/covers/algebra.jpg",\n'
    '    "price": "25000.00",\n'
    '    "status": "active",\n'
    '    "is_offline": false,\n'
    '    "quantity": null,\n'
    '    "date": "2025-05-01",\n'
    '    "created_at": "2025-05-12T10:00:00Z",\n'
    '    "updated_at": "2025-05-12T10:00:00Z"\n'
    '  }\n'
    ']'
)

# ══════════════════════════════════════════════════════
#  GET DETAIL
# ══════════════════════════════════════════════════════
h1('2. Bitta kitob detali')
endpoint_line('GET', '/api/v1/book/books/1/')
lbl('Response  200 OK — yuqoridagi bitta object')
lbl('Xato — 404:')
code_block('{ "detail": "Kitob topilmadi." }')

# ══════════════════════════════════════════════════════
#  POST — ikkala holat
# ══════════════════════════════════════════════════════
h1('3. Kitob yaratish')
endpoint_line('POST', '/api/v1/book/books/', 'faqat superadmin')
info('Content-Type: multipart/form-data — fayl yuklash bo\'lsa shart, aks holda application/json ham bo\'ladi')

lbl('Barcha fieldlar:')
tbl(
    ['Field', 'Turi', 'Majburiy', 'Izoh'],
    [
        ['name',        'string',   'Ha',     'Kitob nomi'],
        ['description', 'string',   'Yo\'q',  'Tavsif'],
        ['category',    'integer',  'Yo\'q',  'Kategoriya ID raqami'],
        ['tags',        'integer[]','Yo\'q',  'Teg ID lari — har biri alohida field sifatida'],
        ['file',        'file',     'Yo\'q',  'PDF yoki boshqa fayl (online kitob uchun)'],
        ['cover_image', 'image',    'Yo\'q',  'JPG / PNG muqova rasmi'],
        ['price',       'number',   'Ha',     'Narx (so\'m)'],
        ['status',      'string',   'Ha',     'active  |  inactive  |  draft'],
        ['is_offline',  'boolean',  'Yo\'q',  'true = oflayn kitob | false = online PDF'],
        ['quantity',    'integer',  'Yo\'q',  'Ombordagi soni — oflayn kitob uchun'],
        ['date',        'date',     'Ha',     'Format: YYYY-MM-DD'],
    ]
)

h2('Holat 1 — Online PDF kitob (application/json):')
code_block(
    '{\n'
    '  "name": "Algebra asoslari",\n'
    '  "description": "PDF qo\'llanma",\n'
    '  "category": 1,\n'
    '  "tags": [1, 2],\n'
    '  "price": 25000,\n'
    '  "status": "active",\n'
    '  "is_offline": false,\n'
    '  "quantity": null,\n'
    '  "date": "2025-05-01"\n'
    '}'
)
info('Fayl yuklash bo\'lsa — multipart/form-data ishlating va file field qo\'shing')

h2('Holat 2 — Oflayn (jismoniy) kitob (multipart/form-data):')
code_block(
    'name          →  Algebra (qog\'oz)\n'
    'description   →  Jismoniy kitob\n'
    'category      →  1\n'
    'tags          →  1\n'
    'tags          →  2\n'
    'price         →  35000\n'
    'status        →  active\n'
    'is_offline    →  true\n'
    'quantity      →  50\n'
    'date          →  2025-05-01\n'
    'cover_image   →  [muqova.jpg]'
)
info('is_offline=true bo\'lsa file yuborilmasa ham bo\'ladi')

lbl('Response  201 Created — to\'liq kitob objecti:')
code_block(
    '{\n'
    '  "id": 5,\n'
    '  "name": "Algebra (qog\'oz)",\n'
    '  "description": "Jismoniy kitob",\n'
    '  "category": { "id": 1, "name": "Matematika", ... },\n'
    '  "tags": [ { "id": 1, "name": "beginner", ... } ],\n'
    '  "file": null,\n'
    '  "cover_image": "https://api.iqmath.uz/media/books/covers/muqova.jpg",\n'
    '  "price": "35000.00",\n'
    '  "status": "active",\n'
    '  "is_offline": true,\n'
    '  "quantity": 50,\n'
    '  "date": "2025-05-01",\n'
    '  "created_at": "...",\n'
    '  "updated_at": "..."\n'
    '}'
)

# ══════════════════════════════════════════════════════
#  PUT — ikkala holat
# ══════════════════════════════════════════════════════
h1('4. Kitobni tahrirlash (Partial update)')
endpoint_line('PUT', '/api/v1/book/books/1/', 'faqat superadmin')
info('Faqat o\'zgartiriladigan fieldlarni yuboring — qolganlar o\'zgarmaydi')
info('Fayl o\'zgartirilsa: multipart/form-data  |  Faqat matn o\'zgarsa: application/json ham bo\'ladi')

h2('Misol 1 — Narx va statusni o\'zgartirish (application/json):')
code_block(
    '{\n'
    '  "price": "30000.00",\n'
    '  "status": "inactive"\n'
    '}'
)

h2('Misol 2 — Online kitobni oflayn ga o\'tkazish (application/json):')
code_block(
    '{\n'
    '  "is_offline": true,\n'
    '  "quantity": 30\n'
    '}'
)

h2('Misol 3 — Ombor sonini yangilash (application/json):')
code_block(
    '{\n'
    '  "quantity": 100\n'
    '}'
)

h2('Misol 4 — Fayl almashtirish (multipart/form-data):')
code_block(
    'file          →  [yangi_algebra.pdf]\n'
    'cover_image   →  [yangi_muqova.png]'
)

lbl('Response  200 OK — yangilangan to\'liq kitob objecti')
lbl('Xato — 404:')
code_block('{ "detail": "Kitob topilmadi." }')

# ══════════════════════════════════════════════════════
#  DELETE
# ══════════════════════════════════════════════════════
h1('5. Kitobni o\'chirish')
endpoint_line('DELETE', '/api/v1/book/books/1/', 'faqat superadmin')
lbl('Response  204 No Content:')
code_block('{ "detail": "Kitob o\'chirildi." }')
lbl('Xato — 404:')
code_block('{ "detail": "Kitob topilmadi." }')

# ══════════════════════════════════════════════════════
#  RUXSATLAR VA XATOLAR
# ══════════════════════════════════════════════════════
h1('6. Ruxsatlar va xato kodlari')

lbl('Ruxsatlar:')
tbl(
    ['Method', 'Kim foydalana oladi'],
    [
        ['GET  (list va detail)', 'Barcha tizimga kirgan foydalanuvchilar'],
        ['POST',   'Faqat superadmin'],
        ['PUT',    'Faqat superadmin'],
        ['DELETE', 'Faqat superadmin'],
    ]
)

lbl('Xato kodlari:')
tbl(
    ['Status', 'Sabab'],
    [
        ['400 Bad Request',  'Majburiy field yo\'q yoki noto\'g\'ri format'],
        ['401 Unauthorized', 'Token yuborilmagan yoki muddati tugagan'],
        ['403 Forbidden',    'Foydalanuvchi superadmin emas'],
        ['404 Not Found',    'Berilgan ID bo\'yicha kitob topilmadi'],
    ]
)

doc.save(r'D:\FASTAPI\iq_math_orginal\api.docx')
print('OK: api.docx yaratildi')
